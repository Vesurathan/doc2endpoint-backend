from docx import Document
import re


def _safe_col_name(name: str) -> str:
    name = str(name).strip().lower()
    name = re.sub(r"[^a-z0-9]+", "_", name)
    return re.sub(r"_+", "_", name).strip("_") or "column"


def extract(file_path: str) -> dict:
    doc = Document(file_path)

    if doc.tables:
        best = max(doc.tables, key=lambda t: len(t.rows))
        if len(best.rows) < 2:
            return _text_fallback(doc)

        headers = [
            cell.text.strip() or f"col_{i}"
            for i, cell in enumerate(best.rows[0].cells)
        ]
        columns = []
        for i, header in enumerate(headers):
            sample = [
                best.rows[r].cells[i].text.strip()
                for r in range(1, min(6, len(best.rows)))
                if i < len(best.rows[r].cells)
            ]
            columns.append({
                "name":          _safe_col_name(header),
                "original_name": header,
                "type":          "string",
                "description":   "",
                "expose":        True,
                "sample_values": [s for s in sample if s],
                "nullable":      True,
            })

        # Build data_preview from actual table rows
        data_preview = []
        for r in range(1, min(6, len(best.rows))):
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(best.rows[r].cells):
                    row_dict[header] = best.rows[r].cells[i].text.strip()
                else:
                    row_dict[header] = ""
            data_preview.append(row_dict)

        return {
            "columns":      columns,
            "row_count":    len(best.rows) - 1,
            "data_preview": data_preview,
        }

    return _text_fallback(doc)


def _text_fallback(doc: Document) -> dict:
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Grab full text excerpt for agent context
    full_text = "\n".join(paragraphs)[:3000]
    return {
        "columns": [{
            "name":          "content",
            "original_name": "content",
            "type":          "string",
            "description":   "Paragraph text",
            "expose":        True,
            "sample_values": paragraphs[:5],
            "nullable":      False,
        }],
        "row_count":    len(paragraphs),
        "data_preview": [{"content": p} for p in paragraphs[:5]],
        "raw_text":     full_text,
        "note":         "No structured tables found. Extracted paragraph text.",
    }
