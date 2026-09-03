import os
import re
import uuid
from datetime import datetime, timezone

import aiofiles
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Request, Query
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.orm import Session

from core.config import settings
from db.database import get_db, engine
from models.dataset import Dataset, DatasetStatus, DocType, ConversationMessage
from models.user import User
from routes.auth import get_current_user
from extractors.dispatcher import extract as do_extract
from agent import schema_agent
from agent import heuristics
from models.user import Plan
from services import storage

router = APIRouter(prefix="/datasets", tags=["datasets"])

ALLOWED_EXTENSIONS = {
    "excel": [".xlsx", ".xls"],
    "csv": [".csv"],
    "pdf": [".pdf"],
    "docx": [".docx", ".doc"],
    "image": [".jpg", ".jpeg", ".png", ".webp", ".tiff"],
}
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _safe_table_name(base: str, user_id: int) -> str:
    base = re.sub(r"[^a-z0-9]+", "_", base.lower()).strip("_")[:40]
    suffix = uuid.uuid4().hex[:8]
    return f"ds_{user_id}_{base}_{suffix}"


def _col_sql_type(col_type: str) -> str:
    return {
        "integer": "BIGINT",
        "number": "DOUBLE PRECISION",
        "boolean": "BOOLEAN",
        "datetime": "TIMESTAMP WITH TIME ZONE",
    }.get(col_type, "TEXT")


# ─── Upload ───────────────────────────────────────────────────────────────────

@router.post("/upload", status_code=201)
async def upload_dataset(
    file: UploadFile = File(...),
    doc_type: str = Form(...),
    name: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if doc_type not in DocType.__members__:
        raise HTTPException(400, "Invalid document type")

    ext = os.path.splitext(file.filename or "")[1].lower()
    allowed = ALLOWED_EXTENSIONS.get(doc_type, [])
    if ext not in allowed:
        raise HTTPException(400, f"File type {ext} not allowed for {doc_type}. Expected: {allowed}")

    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    filename = f"{uuid.uuid4().hex}{ext}"
    local_path = os.path.join(settings.UPLOAD_DIR, filename)

    size = 0
    async with aiofiles.open(local_path, "wb") as out:
        while chunk := await file.read(1024 * 256):
            size += len(chunk)
            if size > MAX_FILE_SIZE:
                os.remove(local_path)
                raise HTTPException(413, "File exceeds 100MB limit")
            await out.write(chunk)

    # Upload to R2 if configured; fall back to local path
    r2_key = f"uploads/{current_user.id}/{filename}"
    stored_path = storage.upload_file(local_path, r2_key)
    if storage._use_r2():
        os.remove(local_path)

    dataset = Dataset(
        user_id=current_user.id,
        name=name or os.path.splitext(file.filename or "untitled")[0],
        doc_type=DocType(doc_type),
        original_filename=file.filename or filename,
        file_path=stored_path,
        status=DatasetStatus.extracting,
    )
    db.add(dataset)
    db.commit()
    db.refresh(dataset)

    # Get a local path for extraction (downloads from R2 if needed)
    extract_path = storage.download_to_temp(stored_path, suffix=ext)
    try:
        schema = do_extract(extract_path, DocType(doc_type))
        dataset.row_count = schema.get("row_count", 0)
        dataset.status = DatasetStatus.reviewing

        # ── Rule-based heuristic analysis (free, instant) ───────────────────
        analysis = heuristics.analyze(schema)
        enriched_cols = analysis["columns"]
        ai_mode = analysis["ai_mode"]

        # ── Override rules — cases that NEVER need AI ────────────────────────
        # Excel / CSV are already fully structured — column names are the headers
        # the user defined; no ambiguity to resolve.
        if DocType(doc_type) in (DocType.excel, DocType.csv):
            ai_mode = "editor"

        # Plain-text documents produce a single generic "content" column.
        # There is nothing for AI to rename/retype — go straight to editor.
        if len(enriched_cols) == 1 and enriched_cols[0].get("name") in ("content", "text"):
            ai_mode = "editor"

        # Free plan users always use the editor (no AI cost)
        if current_user.plan == Plan.free:
            ai_mode = "editor"

        schema["columns"]   = enriched_cols
        schema["ai_mode"]   = ai_mode
        schema["confidence"] = analysis["confidence"]

        dataset.extracted_schema = schema
        dataset.confirmed_schema = {"columns": enriched_cols, "dataset_name": dataset.name}

        # ── Only call AI when genuinely needed ───────────────────────────────
        # (PDF / DOCX with multi-column tables whose names are ambiguous)
        if ai_mode == "chat":
            try:
                opening, initial_cols = schema_agent.get_opening_message(schema)
                schema["columns"] = initial_cols
                dataset.extracted_schema = schema
                dataset.confirmed_schema = {"columns": initial_cols, "dataset_name": dataset.name}
                db.add(ConversationMessage(dataset_id=dataset.id, role="assistant", content=opening))
            except Exception as ai_err:
                # AI unavailable (bad key, network, quota) — degrade gracefully
                # to editor mode so the upload still succeeds.
                ai_mode = "editor"
                schema["ai_mode"] = "editor"
                schema["ai_fallback_note"] = (
                    "AI assistant temporarily unavailable — schema editor shown instead. "
                    f"({type(ai_err).__name__}: {str(ai_err)[:120]})"
                )
                dataset.extracted_schema = schema

        db.commit()
        db.refresh(dataset)
    except Exception as e:
        dataset.status = DatasetStatus.failed
        db.commit()
        raise HTTPException(500, f"Extraction failed: {str(e)}")
    finally:
        # Clean up the temp extraction file when using R2
        if storage._use_r2() and extract_path != stored_path:
            try:
                os.remove(extract_path)
            except OSError:
                pass

    return _dataset_out(dataset)


# ─── List + Get ───────────────────────────────────────────────────────────────

@router.get("/")
def list_datasets(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    datasets = db.query(Dataset).filter(Dataset.user_id == current_user.id).order_by(Dataset.created_at.desc()).all()
    return [_dataset_out(d) for d in datasets]


@router.get("/{dataset_id}")
def get_dataset(dataset_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    dataset = _get_or_404(dataset_id, current_user.id, db)
    return _dataset_out(dataset)


# ─── Chat ─────────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str


@router.get("/{dataset_id}/messages")
def get_messages(dataset_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    dataset = _get_or_404(dataset_id, current_user.id, db)
    return [{"role": m.role, "content": m.content, "id": m.id} for m in dataset.messages]


@router.post("/{dataset_id}/chat")
def chat(dataset_id: int, body: ChatRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    # Plan gate — free users must use the schema editor
    if current_user.plan == Plan.free:
        raise HTTPException(
            403,
            "AI chat is available on Premium and Pro plans. "
            "Use the schema editor to review your columns, or upgrade for AI assistance."
        )

    dataset = _get_or_404(dataset_id, current_user.id, db)
    if dataset.status not in (DatasetStatus.reviewing,):
        raise HTTPException(400, "Dataset is not in review state")

    # Cap history to last 6 messages (3 turns) — enough context, cuts token cost
    all_history = [{"role": m.role, "content": m.content} for m in dataset.messages]
    history = all_history[-6:] if len(all_history) > 6 else all_history

    current_cols = (dataset.confirmed_schema or {}).get("columns", dataset.extracted_schema.get("columns", []))

    reply, updated_cols = schema_agent.chat(
        history=history,
        extracted_schema=dataset.extracted_schema,
        current_columns=current_cols,
        user_message=body.message,
    )

    db.add(ConversationMessage(dataset_id=dataset.id, role="user", content=body.message))
    db.add(ConversationMessage(dataset_id=dataset.id, role="assistant", content=reply))

    confirmed = dataset.confirmed_schema or {}
    confirmed["columns"] = updated_cols
    dataset.confirmed_schema = confirmed
    dataset.updated_at = datetime.now(timezone.utc)
    db.commit()

    return {"reply": reply, "updated_columns": updated_cols}


# ─── Confirm → Create Table + Load Data ──────────────────────────────────────

class ConfirmRequest(BaseModel):
    columns: list[dict] | None = None  # supplied by the schema editor


@router.post("/{dataset_id}/confirm")
def confirm_schema(dataset_id: int, body: ConfirmRequest = ConfirmRequest(), db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    dataset = _get_or_404(dataset_id, current_user.id, db)
    if dataset.status != DatasetStatus.reviewing:
        raise HTTPException(400, "Dataset must be in reviewing state to confirm")

    # Schema editor sends columns directly; AI chat path uses what's already stored
    if body.columns:
        dataset.confirmed_schema = {"columns": body.columns, "dataset_name": dataset.name}

    confirmed = dataset.confirmed_schema or {}
    columns = confirmed.get("columns", [])
    if not columns:
        raise HTTPException(400, "No confirmed schema found")

    exposed = [c for c in columns if c.get("expose", True)]
    if not exposed:
        raise HTTPException(400, "At least one column must be exposed")

    table_name = _safe_table_name(dataset.name, current_user.id)

    col_defs = ", ".join(
        f'"{c["name"]}" {_col_sql_type(c.get("type", "string"))}'
        for c in exposed
    )
    create_sql = f'CREATE TABLE IF NOT EXISTS "{table_name}" (id SERIAL PRIMARY KEY, {col_defs})'

    with engine.begin() as conn:
        conn.execute(text(create_sql))
        _load_data(conn, table_name, exposed, dataset)

    dataset.table_name = table_name
    dataset.status = DatasetStatus.active
    dataset.confirmed_schema = confirmed
    db.commit()
    db.refresh(dataset)

    return {**_dataset_out(dataset), "api_base": f"/api/v1/{table_name}"}


def _load_data(conn, table_name: str, columns: list[dict], dataset: Dataset):
    import pandas as pd

    doc_type = dataset.doc_type
    ext = os.path.splitext(dataset.original_filename or "")[1].lower() or ""
    fp = storage.download_to_temp(dataset.file_path, suffix=ext)
    _is_temp = storage._use_r2() and fp != dataset.file_path

    try:
        # ── Structured formats: Excel / CSV → pandas ─────────────────────────
        if doc_type in (DocType.excel,):
            df = pd.read_excel(fp)
        elif doc_type == DocType.csv:
            df = pd.read_csv(fp)

        # ── PDF ───────────────────────────────────────────────────────────────
        elif doc_type == DocType.pdf:
            _load_pdf(conn, table_name, columns, fp)
            return

        # ── DOCX ──────────────────────────────────────────────────────────────
        elif doc_type == DocType.docx:
            _load_docx(conn, table_name, columns, fp)
            return

        else:
            return  # image / unsupported — nothing to load
    finally:
        if _is_temp:
            try:
                os.remove(fp)
            except OSError:
                pass

    # ── Common path for Excel / CSV ───────────────────────────────────────────
    df = df.dropna(how="all")

    # Map original column names (as they appear in the file) → API-safe names
    col_name_map = {c["original_name"]: c["name"] for c in columns}

    # Keep only columns that exist in the file AND are in the confirmed schema
    existing_originals = [orig for orig in col_name_map if orig in df.columns]
    if not existing_originals:
        return

    # Select and rename in one step — df now has API-safe column names
    df = df[existing_originals].rename(columns=col_name_map)

    if df.empty:
        return

    placeholders = ", ".join(f":{c}" for c in df.columns)
    col_list = ", ".join(f'"{c}"' for c in df.columns)
    insert_sql = text(f'INSERT INTO "{table_name}" ({col_list}) VALUES ({placeholders})')

    records = df.where(df.notna(), None).to_dict(orient="records")
    conn.execute(insert_sql, records)


def _load_pdf(conn, table_name: str, columns: list[dict], file_path: str):
    """Load PDF data into the database table.

    Two cases:
    • PDF with tables  → re-extract with pdfplumber, map header→api_name, insert all rows.
    • PDF without tables → split raw text into lines, insert each as a 'content' row.
    """
    import pdfplumber
    import re

    def _safe(name: str) -> str:
        name = str(name).strip().lower()
        name = re.sub(r"[^a-z0-9]+", "_", name)
        return re.sub(r"_+", "_", name).strip("_") or "column"

    # Build original_name → api_name map from confirmed schema
    col_map = {c["original_name"]: c["name"] for c in columns}
    exposed_names = {c["name"] for c in columns}

    all_rows: list[dict] = []
    tables_found = False

    with pdfplumber.open(file_path) as pdf:
        page_tables = []
        for page in pdf.pages:
            for tbl in page.extract_tables():
                if tbl and len(tbl) >= 2:
                    page_tables.append(tbl)

        if page_tables:
            tables_found = True
            # Use the largest table (same heuristic as extractor)
            best = max(page_tables, key=lambda t: len(t))
            raw_headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(best[0])]
            # Map extracted header → api_name via col_map (original_name key)
            header_to_api = {h: col_map.get(h, _safe(h)) for h in raw_headers}

            for row in best[1:]:
                record = {}
                for i, header in enumerate(raw_headers):
                    api_name = header_to_api[header]
                    if api_name in exposed_names:
                        record[api_name] = str(row[i]).strip() if i < len(row) and row[i] is not None else None
                if record:
                    all_rows.append(record)

        if not tables_found:
            # Plain-text PDF — collect all lines, store in 'content' column
            content_col = next((c["name"] for c in columns if c["name"] == "content"), None)
            if not content_col and columns:
                content_col = columns[0]["name"]
            if content_col:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    for line in text.splitlines():
                        line = line.strip()
                        if line:
                            all_rows.append({content_col: line})

    if not all_rows:
        return

    # Use the first row's keys to build the INSERT
    first = all_rows[0]
    col_list = ", ".join(f'"{k}"' for k in first)
    placeholders = ", ".join(f":{k}" for k in first)
    insert_sql = text(f'INSERT INTO "{table_name}" ({col_list}) VALUES ({placeholders})')
    conn.execute(insert_sql, all_rows)


def _load_docx(conn, table_name: str, columns: list[dict], file_path: str):
    """Load DOCX data into the database table.

    Two cases:
    • DOCX with tables  → iterate all rows of the largest table, map header→api_name.
    • DOCX without tables → insert each paragraph as a 'content' row.
    """
    from docx import Document

    col_map = {c["original_name"]: c["name"] for c in columns}
    exposed_names = {c["name"] for c in columns}

    doc = Document(file_path)
    all_rows: list[dict] = []

    if doc.tables:
        best = max(doc.tables, key=lambda t: len(t.rows))
        if len(best.rows) >= 2:
            raw_headers = [cell.text.strip() or f"col_{i}" for i, cell in enumerate(best.rows[0].cells)]
            header_to_api = {h: col_map.get(h, h) for h in raw_headers}

            for row in best.rows[1:]:
                record = {}
                for i, header in enumerate(raw_headers):
                    api_name = header_to_api[header]
                    if api_name in exposed_names:
                        record[api_name] = row.cells[i].text.strip() if i < len(row.cells) else None
                if record:
                    all_rows.append(record)

    if not all_rows:
        # Paragraph fallback
        content_col = next((c["name"] for c in columns if c["name"] == "content"), None)
        if not content_col and columns:
            content_col = columns[0]["name"]
        if content_col:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_rows.append({content_col: text})

    if not all_rows:
        return

    first = all_rows[0]
    col_list = ", ".join(f'"{k}"' for k in first)
    placeholders = ", ".join(f":{k}" for k in first)
    insert_sql = text(f'INSERT INTO "{table_name}" ({col_list}) VALUES ({placeholders})')
    conn.execute(insert_sql, all_rows)


# ─── Preview (playground — JWT auth instead of API key) ──────────────────────

@router.get("/{dataset_id}/preview")
def preview_dataset(
    dataset_id: int,
    request: Request,
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    order_by: str = Query("id"),
    order: str = Query("asc"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    dataset = _get_or_404(dataset_id, current_user.id, db)
    if dataset.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active to preview")

    cols = dataset.confirmed_schema.get("columns", [])
    valid = {c["name"] for c in cols if c.get("expose", True)} | {"id"}

    filter_params = {
        k: v for k, v in request.query_params.items()
        if k not in ("page", "limit", "order_by", "order")
    }

    conditions, bind = [], {}
    for key, val in filter_params.items():
        if key.endswith("__gte"):
            col = key[:-5]
            if col in valid:
                conditions.append(f'"{col}" >= :gte_{col}')
                bind[f"gte_{col}"] = val
        elif key.endswith("__lte"):
            col = key[:-5]
            if col in valid:
                conditions.append(f'"{col}" <= :lte_{col}')
                bind[f"lte_{col}"] = val
        elif key.endswith("__like"):
            col = key[:-6]
            if col in valid:
                conditions.append(f'"{col}"::TEXT ILIKE :like_{col}')
                bind[f"like_{col}"] = f"%{val}%"
        elif key in valid:
            conditions.append(f'"{key}" = :eq_{key}')
            bind[f"eq_{key}"] = val

    where = f"WHERE {' AND '.join(conditions)}" if conditions else ""
    sort_col = order_by if order_by in valid else "id"
    sort_dir = "DESC" if order.upper() == "DESC" else "ASC"
    offset = (page - 1) * limit

    try:
        with engine.connect() as conn:
            total = conn.execute(text(f'SELECT COUNT(*) FROM "{dataset.table_name}" {where}'), bind).scalar()
            bind["limit"] = limit
            bind["offset"] = offset
            rows = conn.execute(
                text(f'SELECT * FROM "{dataset.table_name}" {where} ORDER BY "{sort_col}" {sort_dir} LIMIT :limit OFFSET :offset'),
                bind,
            ).mappings().all()
    except Exception as e:
        raise HTTPException(500, f"Query error: {str(e)}")

    return {
        "data": [dict(r) for r in rows],
        "total": total,
        "page": page,
        "limit": limit,
        "pages": max(1, -(-total // limit)),
        "columns": [c["name"] for c in cols if c.get("expose", True)],
    }


# ─── OpenAPI spec ─────────────────────────────────────────────────────────────

@router.get("/{dataset_id}/openapi")
def get_openapi_spec(
    dataset_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    dataset = _get_or_404(dataset_id, current_user.id, db)
    if dataset.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active to generate OpenAPI spec")

    cols = dataset.confirmed_schema.get("columns", [])
    exposed = [c for c in cols if c.get("expose", True)]
    table = dataset.table_name
    schema_name = re.sub(r"[^a-zA-Z0-9]", "", dataset.name) or "Dataset"

    _TYPE = {"string": "string", "integer": "integer", "number": "number", "boolean": "boolean", "datetime": "string"}
    _FMT  = {"datetime": "date-time"}

    properties = {"id": {"type": "integer", "description": "Auto-generated record ID"}}
    for col in exposed:
        ct = col.get("type", "string")
        prop: dict = {"type": _TYPE.get(ct, "string")}
        if ct in _FMT:
            prop["format"] = _FMT[ct]
        if col.get("description"):
            prop["description"] = col["description"]
        properties[col["name"]] = prop

    filter_params = [
        {"name": "page",     "in": "query", "description": "Page number (1-indexed)",   "schema": {"type": "integer", "default": 1,   "minimum": 1}},
        {"name": "limit",    "in": "query", "description": "Records per page (max 1000)", "schema": {"type": "integer", "default": 100, "maximum": 1000}},
        {"name": "order_by", "in": "query", "description": "Column name to sort by",    "schema": {"type": "string"}},
        {"name": "order",    "in": "query", "description": "Sort direction",             "schema": {"type": "string", "enum": ["asc", "desc"], "default": "asc"}},
    ]
    for col in exposed:
        cn, ct = col["name"], col.get("type", "string")
        ot = _TYPE.get(ct, "string")
        filter_params.append({"name": cn,              "in": "query", "description": f"Exact match on {cn}",        "schema": {"type": ot}})
        if ct == "string":
            filter_params.append({"name": f"{cn}__like", "in": "query", "description": f"{cn} contains (case-insensitive)", "schema": {"type": "string"}})
        else:
            filter_params.append({"name": f"{cn}__gte", "in": "query", "description": f"{cn} ≥ value",        "schema": {"type": ot}})
            filter_params.append({"name": f"{cn}__lte", "in": "query", "description": f"{cn} ≤ value",        "schema": {"type": ot}})

    spec = {
        "openapi": "3.0.3",
        "info": {
            "title": dataset.name,
            "description": f"Auto-generated REST API for the **{dataset.name}** dataset.\n\nCreated with [Doc2Endpoint](https://doc2endpoint.io). **{dataset.row_count:,} rows.**",
            "version": "1.0.0",
        },
        "servers": [
            {"url": "http://localhost:8000",   "description": "Local development"},
            {"url": "https://api.doc2endpoint.io",   "description": "Production"},
        ],
        "security": [{"ApiKeyAuth": []}],
        "components": {
            "securitySchemes": {
                "ApiKeyAuth": {"type": "apiKey", "in": "header", "name": "X-API-Key", "description": "Your Doc2Endpoint API key"},
            },
            "schemas": {
                schema_name: {"type": "object", "properties": properties},
                "PaginatedResponse": {
                    "type": "object",
                    "properties": {
                        "data":  {"type": "array", "items": {"$ref": f"#/components/schemas/{schema_name}"}},
                        "total": {"type": "integer"},
                        "page":  {"type": "integer"},
                        "pages": {"type": "integer"},
                        "limit": {"type": "integer"},
                    },
                },
            },
        },
        "paths": {
            f"/api/v1/{table}": {
                "get": {
                    "summary": f"List {dataset.name} records",
                    "description": f"Returns paginated, filterable records from the **{dataset.name}** dataset.",
                    "operationId": f"list_{table}",
                    "parameters": filter_params,
                    "responses": {
                        "200": {"description": "Paginated list of records", "content": {"application/json": {"schema": {"$ref": "#/components/schemas/PaginatedResponse"}}}},
                        "401": {"description": "Missing API key"},
                        "403": {"description": "Invalid or unauthorised API key"},
                    },
                },
            },
            f"/api/v1/{table}/schema": {
                "get": {
                    "summary": "Get dataset schema",
                    "operationId": f"schema_{table}",
                    "responses": {"200": {"description": "Column definitions and row count"}},
                },
            },
            f"/api/v1/{table}/{{id}}": {
                "get": {
                    "summary": f"Get single {dataset.name} record",
                    "operationId": f"get_{table}",
                    "parameters": [{"name": "id", "in": "path", "required": True, "schema": {"type": "integer"}}],
                    "responses": {
                        "200": {"description": "Single record", "content": {"application/json": {"schema": {"type": "object", "properties": {"data": {"$ref": f"#/components/schemas/{schema_name}"}}}}}},
                        "404": {"description": "Record not found"},
                    },
                },
            },
        },
    }

    return spec


# ─── Public OpenAPI spec (no auth — for GPT Actions / Zapier) ────────────────

@router.get("/{dataset_id}/openapi/public")
def get_openapi_spec_public(
    dataset_id: int,
    db: Session = Depends(get_db),
):
    """
    Public version of the OpenAPI spec endpoint — no JWT required.
    Used by ChatGPT Actions, Zapier, and other external tooling to fetch
    the schema. Returns 404 if the dataset does not exist or is not active.
    The spec itself describes query parameters and response shapes; actual
    data access still requires a valid X-API-Key.
    """
    ds = db.query(Dataset).filter(
        Dataset.id == dataset_id,
        Dataset.status == DatasetStatus.active,
    ).first()
    if not ds:
        raise HTTPException(404, "Dataset not found or not yet active")

    # Re-use the authenticated spec builder logic inline
    from routes.datasets import get_openapi_spec
    # Build a minimal user stub — the spec builder only needs ds fields
    cols = ds.confirmed_schema.get("columns", [])
    exposed = [c for c in cols if c.get("expose", True)]
    table = ds.table_name
    schema_name = re.sub(r"[^a-zA-Z0-9]", "", ds.name) or "Dataset"

    _TYPE = {"string": "string", "integer": "integer", "number": "number", "boolean": "boolean", "datetime": "string"}
    _FMT  = {"datetime": "date-time"}

    properties = {"id": {"type": "integer", "description": "Auto-generated record ID"}}
    for col in exposed:
        ct = col.get("type", "string")
        prop: dict = {"type": _TYPE.get(ct, "string")}
        if ct in _FMT:
            prop["format"] = _FMT[ct]
        if col.get("description"):
            prop["description"] = col["description"]
        properties[col["name"]] = prop

    filter_params = [
        {"name": "page",     "in": "query", "schema": {"type": "integer", "default": 1}},
        {"name": "limit",    "in": "query", "schema": {"type": "integer", "default": 100, "maximum": 1000}},
        {"name": "order_by", "in": "query", "schema": {"type": "string"}},
        {"name": "order",    "in": "query", "schema": {"type": "string", "enum": ["asc", "desc"]}},
    ]
    for col in exposed:
        cn, ct = col["name"], col.get("type", "string")
        ot = _TYPE.get(ct, "string")
        filter_params.append({"name": cn, "in": "query", "schema": {"type": ot}})

    api_base = settings.FRONTEND_URL.replace("5173", "8000") if "5173" in settings.FRONTEND_URL else settings.FRONTEND_URL

    return {
        "openapi": "3.0.3",
        "info": {
            "title": ds.name,
            "description": f"Auto-generated REST API for **{ds.name}**. Created with [Doc2Endpoint](https://doc2endpoint.io). {ds.row_count:,} rows.",
            "version": "1.0.0",
        },
        "servers": [{"url": api_base, "description": "Doc2Endpoint server"}],
        "security": [{"ApiKeyAuth": []}],
        "components": {
            "securitySchemes": {
                "ApiKeyAuth": {"type": "apiKey", "in": "header", "name": "X-API-Key"},
            },
            "schemas": {
                schema_name: {"type": "object", "properties": properties},
                "PaginatedResponse": {
                    "type": "object",
                    "properties": {
                        "data":  {"type": "array", "items": {"$ref": f"#/components/schemas/{schema_name}"}},
                        "total": {"type": "integer"},
                        "page":  {"type": "integer"},
                        "pages": {"type": "integer"},
                        "limit": {"type": "integer"},
                    },
                },
            },
        },
        "paths": {
            f"/api/v1/{table}": {
                "get": {
                    "summary": f"List {ds.name} records",
                    "operationId": f"list_{table}",
                    "parameters": filter_params,
                    "responses": {
                        "200": {"description": "Paginated records", "content": {"application/json": {"schema": {"$ref": "#/components/schemas/PaginatedResponse"}}}},
                        "401": {"description": "Missing API key"},
                        "403": {"description": "Invalid API key"},
                    },
                },
            },
            f"/api/v1/{table}/{{id}}": {
                "get": {
                    "summary": f"Get single {ds.name} record",
                    "operationId": f"get_{table}",
                    "parameters": [{"name": "id", "in": "path", "required": True, "schema": {"type": "integer"}}],
                    "responses": {
                        "200": {"description": "Single record"},
                        "404": {"description": "Record not found"},
                    },
                },
            },
        },
    }


# ─── Activity log ─────────────────────────────────────────────────────────────

@router.get("/{dataset_id}/logs")
def get_dataset_logs(
    dataset_id: int,
    key_id: int | None = Query(None),
    status: int | None = Query(None),
    from_date: str | None = Query(None, alias="from"),
    to_date: str | None = Query(None, alias="to"),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from models.api_usage import APIUsage
    from models.api_key import APIKey as APIKeyModel
    from sqlalchemy import and_

    dataset = _get_or_404(dataset_id, current_user.id, db)

    q = db.query(APIUsage).filter(APIUsage.dataset_id == dataset.id)

    if key_id is not None:
        q = q.filter(APIUsage.api_key_id == key_id)
    if status is not None:
        q = q.filter(APIUsage.status_code == status)
    if from_date:
        try:
            q = q.filter(APIUsage.created_at >= datetime.fromisoformat(from_date))
        except ValueError:
            pass
    if to_date:
        try:
            q = q.filter(APIUsage.created_at <= datetime.fromisoformat(to_date))
        except ValueError:
            pass

    total = q.count()
    logs = q.order_by(APIUsage.created_at.desc()).offset(offset).limit(limit).all()

    # Build a key_id → key_prefix lookup to show friendly names
    key_ids = {log.api_key_id for log in logs}
    keys = {k.id: k for k in db.query(APIKeyModel).filter(APIKeyModel.id.in_(key_ids)).all()} if key_ids else {}

    return {
        "total": total,
        "offset": offset,
        "limit": limit,
        "logs": [
            {
                "id":           log.id,
                "method":       log.method,
                "path":         log.path,
                "query_string": log.query_string,
                "status_code":  log.status_code,
                "response_ms":  log.response_time_ms,
                "ip_address":   log.ip_address,
                "created_at":   log.created_at.isoformat(),
                "api_key_id":   log.api_key_id,
                "api_key_prefix": keys[log.api_key_id].key_prefix if log.api_key_id in keys else None,
                "api_key_name": keys[log.api_key_id].name if log.api_key_id in keys else None,
            }
            for log in logs
        ],
    }


# ─── Custom endpoint ──────────────────────────────────────────────────────────

import re as _re

_SLUG_RE = _re.compile(r'^[a-z0-9][a-z0-9_-]{1,63}$')

@router.patch("/{dataset_id}/endpoint")
def set_custom_endpoint(
    dataset_id: int,
    body: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Set or clear a custom URL slug for /api/v1/<slug>.
    Pass {"custom_endpoint": "my-slug"} to set, or {"custom_endpoint": null} to reset to default."""
    dataset = _get_or_404(dataset_id, current_user.id, db)
    if dataset.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active before changing its endpoint.")

    slug = body.get("custom_endpoint")

    if slug is None or slug == "":
        # Clear — revert to table_name
        dataset.custom_endpoint = None
        db.commit()
        return {"custom_endpoint": None, "effective_endpoint": dataset.table_name}

    slug = slug.strip().lower()
    if not _SLUG_RE.match(slug):
        raise HTTPException(
            422,
            "Endpoint slug must be 2–64 characters, start with a letter or digit, "
            "and contain only lowercase letters, digits, hyphens, and underscores."
        )

    # Make sure it doesn't clash with another dataset's table_name or custom_endpoint
    conflict = db.query(Dataset).filter(
        Dataset.id != dataset_id,
        (Dataset.custom_endpoint == slug) | (Dataset.table_name == slug),
    ).first()
    if conflict:
        raise HTTPException(409, f"Endpoint '{slug}' is already in use by another dataset.")

    dataset.custom_endpoint = slug
    db.commit()
    return {"custom_endpoint": slug, "effective_endpoint": slug}


# ─── Shared ───────────────────────────────────────────────────────────────────

def _get_or_404(dataset_id: int, user_id: int, db: Session) -> Dataset:
    ds = db.query(Dataset).filter(Dataset.id == dataset_id, Dataset.user_id == user_id).first()
    if not ds:
        raise HTTPException(404, "Dataset not found")
    return ds


def _dataset_out(d: Dataset) -> dict:
    return {
        "id": d.id,
        "name": d.name,
        "doc_type": d.doc_type,
        "status": d.status,
        "original_filename": d.original_filename,
        "row_count": d.row_count,
        "table_name": d.table_name,
        "extracted_schema": d.extracted_schema,
        "confirmed_schema": d.confirmed_schema,
        "ai_mode": (d.extracted_schema or {}).get("ai_mode", "chat"),
        "created_at": d.created_at.isoformat(),
        # New fields
        "is_public": getattr(d, "is_public", False),
        "public_description": getattr(d, "public_description", None),
        "webhook_url": getattr(d, "webhook_url", None),
        "sync_url": getattr(d, "sync_url", None),
        "sync_interval_hours": getattr(d, "sync_interval_hours", None),
        "last_synced_at": d.last_synced_at.isoformat() if getattr(d, "last_synced_at", None) else None,
        "custom_endpoint": getattr(d, "custom_endpoint", None),
    }


# ─── Public Gallery ───────────────────────────────────────────────────────────

@router.get("/gallery")
def public_gallery(
    search: str = Query(""),
    doc_type: str = Query(""),
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    db: Session = Depends(get_db),
):
    """No auth required — lists all publicly shared datasets."""
    from models.user import User as UserModel
    q = db.query(Dataset).filter(
        Dataset.is_public == True,
        Dataset.status == DatasetStatus.active,
    )
    if search:
        q = q.filter(Dataset.name.ilike(f"%{search}%"))
    if doc_type:
        q = q.filter(Dataset.doc_type == doc_type)

    total = q.count()
    datasets = q.order_by(Dataset.row_count.desc()).offset(offset).limit(limit).all()

    results = []
    for ds in datasets:
        owner = db.query(UserModel).filter(UserModel.id == ds.user_id).first()
        cols = (ds.confirmed_schema or {}).get("columns", [])
        results.append({
            "id": ds.id,
            "name": ds.name,
            "doc_type": ds.doc_type,
            "row_count": ds.row_count,
            "table_name": ds.table_name,
            "public_description": getattr(ds, "public_description", None),
            "columns": [{"name": c["name"], "type": c.get("type", "string")} for c in cols if c.get("expose", True)],
            "created_at": ds.created_at.isoformat(),
            "owner": owner.full_name if owner else "Anonymous",
        })

    return {"data": results, "total": total}


class VisibilityRequest(BaseModel):
    is_public: bool
    public_description: str = ""


@router.patch("/{dataset_id}/visibility")
def set_visibility(
    dataset_id: int,
    body: VisibilityRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    ds = _get_or_404(dataset_id, current_user.id, db)
    if ds.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active to publish to gallery")
    ds.is_public = body.is_public
    ds.public_description = body.public_description.strip() or None
    db.commit()
    db.refresh(ds)
    return _dataset_out(ds)


# ─── Webhooks ─────────────────────────────────────────────────────────────────

class WebhookRequest(BaseModel):
    webhook_url: str = ""
    webhook_secret: str = ""


@router.patch("/{dataset_id}/webhook")
def set_webhook(
    dataset_id: int,
    body: WebhookRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    ds = _get_or_404(dataset_id, current_user.id, db)
    ds.webhook_url = body.webhook_url.strip() or None
    ds.webhook_secret = body.webhook_secret.strip() or None
    db.commit()
    db.refresh(ds)
    return _dataset_out(ds)


@router.post("/{dataset_id}/webhook/test")
def test_webhook(
    dataset_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from services.webhook import fire_webhook
    ds = _get_or_404(dataset_id, current_user.id, db)
    if not ds.webhook_url:
        raise HTTPException(400, "No webhook URL configured")
    ok = fire_webhook(ds.webhook_url, ds.webhook_secret, {
        "event": "dataset.test",
        "dataset_id": ds.id,
        "dataset_name": ds.name,
        "message": "This is a test webhook from Doc2Endpoint.",
    })
    if not ok:
        raise HTTPException(502, "Webhook delivery failed — check the URL and try again")
    return {"message": "Webhook delivered successfully"}


# ─── Scheduled Sync ───────────────────────────────────────────────────────────

class SyncConfigRequest(BaseModel):
    sync_url: str = ""
    sync_interval_hours: int = 24


@router.patch("/{dataset_id}/sync-config")
def set_sync_config(
    dataset_id: int,
    body: SyncConfigRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    ds = _get_or_404(dataset_id, current_user.id, db)
    if ds.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active to configure sync")
    url = body.sync_url.strip()
    if url:
        if body.sync_interval_hours < 1:
            raise HTTPException(400, "sync_interval_hours must be at least 1")
        ds.sync_url = url
        ds.sync_interval_hours = body.sync_interval_hours
    else:
        ds.sync_url = None
        ds.sync_interval_hours = None
    db.commit()
    db.refresh(ds)
    return _dataset_out(ds)


@router.post("/{dataset_id}/sync")
def trigger_sync(
    dataset_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Manual sync trigger."""
    from services.sync import sync_dataset
    ds = _get_or_404(dataset_id, current_user.id, db)
    if ds.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active to sync")
    if not ds.sync_url:
        raise HTTPException(400, "No sync URL configured")
    try:
        row_count = sync_dataset(ds, db)
    except Exception as e:
        raise HTTPException(502, f"Sync failed: {str(e)}")
    return {"message": "Sync completed", "row_count": row_count, "synced_at": ds.last_synced_at.isoformat()}


# ─── Upload New Version ───────────────────────────────────────────────────────

@router.post("/{dataset_id}/upload-version", status_code=200)
async def upload_version(
    dataset_id: int,
    file: UploadFile = File(...),
    mode: str = Form("replace"),   # "replace" | "append"
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload a new version of an existing active dataset."""
    import pandas as pd

    ds = _get_or_404(dataset_id, current_user.id, db)
    if ds.status != DatasetStatus.active:
        raise HTTPException(400, "Dataset must be active to upload a new version")
    if mode not in ("replace", "append"):
        raise HTTPException(400, "mode must be 'replace' or 'append'")

    ext = os.path.splitext(file.filename or "")[1].lower()
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    tmp_path = os.path.join(settings.UPLOAD_DIR, f"{uuid.uuid4().hex}{ext}")

    size = 0
    async with aiofiles.open(tmp_path, "wb") as out:
        while chunk := await file.read(1024 * 256):
            size += len(chunk)
            if size > MAX_FILE_SIZE:
                os.remove(tmp_path)
                raise HTTPException(413, "File exceeds 100MB limit")
            await out.write(chunk)

    try:
        # Parse new file
        if ext in (".xlsx", ".xls"):
            df = pd.read_excel(tmp_path)
        elif ext == ".csv":
            df = pd.read_csv(tmp_path)
        else:
            os.remove(tmp_path)
            raise HTTPException(400, f"Version uploads support CSV and Excel only (got {ext})")

        df = df.dropna(how="all")

        # Map column names to confirmed schema
        cols = (ds.confirmed_schema or {}).get("columns", [])
        exposed = [c for c in cols if c.get("expose", True)]
        col_name_map = {c.get("original_name", c["name"]): c["name"] for c in exposed}

        matched = {orig: new for orig, new in col_name_map.items() if orig in df.columns}
        if not matched:
            matched = {c["name"]: c["name"] for c in exposed if c["name"] in df.columns}

        if not matched:
            os.remove(tmp_path)
            raise HTTPException(400, f"No matching columns. File has: {list(df.columns)}")

        df = df[list(matched.keys())].rename(columns=matched)
        target_cols = [c["name"] for c in exposed if c["name"] in df.columns]
        df = df[target_cols]

        new_rows = len(df)

        with engine.begin() as conn:
            if mode == "replace":
                conn.execute(text(f'TRUNCATE TABLE "{ds.table_name}"'))
            # Insert new rows
            if not df.empty:
                placeholders = ", ".join(f":{c}" for c in df.columns)
                col_list = ", ".join(f'"{c}"' for c in df.columns)
                insert_sql = text(f'INSERT INTO "{ds.table_name}" ({col_list}) VALUES ({placeholders})')
                conn.execute(insert_sql, df.where(df.notna(), None).to_dict(orient="records"))

        # Update row count
        with engine.connect() as conn:
            total = conn.execute(text(f'SELECT COUNT(*) FROM "{ds.table_name}"')).scalar()

        ds.row_count = total
        ds.original_filename = file.filename or ds.original_filename
        ds.updated_at = datetime.now(timezone.utc)
        db.commit()
        db.refresh(ds)

        # Fire webhook if configured
        if ds.webhook_url:
            from services.webhook import fire_webhook
            fire_webhook(ds.webhook_url, ds.webhook_secret, {
                "event": "dataset.version_uploaded",
                "dataset_id": ds.id,
                "dataset_name": ds.name,
                "mode": mode,
                "new_rows": new_rows,
                "total_rows": total,
            })

    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    return {
        **_dataset_out(ds),
        "new_rows": new_rows,
        "mode": mode,
        "message": f"{'Replaced' if mode == 'replace' else 'Appended'} {new_rows:,} rows. Total: {ds.row_count:,}",
    }
